import json
import time
from pathlib import Path

import fitz
import openpyxl
from docx import Document
from typer.testing import CliRunner

from docrt.cli import app

runner = CliRunner()


def test_inspect_docx_output_option_writes_explicit_json(tmp_path: Path):
    input_path = tmp_path / "sample.docx"
    output_path = tmp_path / "custom" / "docx.json"
    document = Document()
    document.add_paragraph("hello docx")
    document.save(input_path)

    result = runner.invoke(app, ["inspect-docx", str(input_path), "--output", str(output_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    saved = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["ok"] is True
    assert payload["output_path"] == str(output_path.resolve())
    assert saved["paragraphs"] == ["hello docx"]


def test_inspect_pdf_output_option_writes_explicit_json(tmp_path: Path):
    input_path = tmp_path / "sample.pdf"
    output_path = tmp_path / "custom" / "pdf.json"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "hello pdf")
    document.save(input_path)
    document.close()

    result = runner.invoke(app, ["inspect-pdf", str(input_path), "--output", str(output_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    saved = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["ok"] is True
    assert payload["output_path"] == str(output_path.resolve())
    assert saved["page_count"] == 1


def test_search_pdf_command_outputs_matches(tmp_path: Path):
    input_path = tmp_path / "sample.pdf"
    output_path = tmp_path / "custom" / "pdf-search.json"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "hello searchable pdf")
    document.save(input_path)
    document.close()

    result = runner.invoke(
        app,
        ["search-pdf", str(input_path), "searchable", "--output", str(output_path)],
    )

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    saved = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["ok"] is True
    assert payload["operation"] == "search-pdf"
    assert payload["data"]["count"] == 1
    assert saved["matches"][0]["page_number"] == 1


def test_search_pdf_command_supports_pages_option(tmp_path: Path):
    input_path = tmp_path / "sample.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "first target")
    page = document.new_page()
    page.insert_text((72, 72), "second target")
    document.save(input_path)
    document.close()

    result = runner.invoke(app, ["search-pdf", str(input_path), "target", "--pages", "2"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["data"]["count"] == 1
    assert payload["data"]["page_selection"]["selected_pages"] == [2]


def test_read_pdf_command_supports_pages_option(tmp_path: Path):
    input_path = tmp_path / "sample.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "first page")
    page = document.new_page()
    page.insert_text((72, 72), "second page")
    document.save(input_path)
    document.close()

    result = runner.invoke(app, ["read-pdf", str(input_path), "--pages", "2"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["data"]["metadata"]["page_selection"]["selected_pages"] == [2]
    assert len(payload["data"]["content_blocks"]) == 1


def test_inspect_xlsx_output_option_writes_explicit_json(tmp_path: Path):
    input_path = tmp_path / "sample.xlsx"
    output_path = tmp_path / "custom" / "xlsx.json"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet["A1"] = "hello xlsx"
    workbook.save(input_path)

    result = runner.invoke(app, ["inspect-xlsx", str(input_path), "--output", str(output_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    saved = json.loads(output_path.read_text(encoding="utf-8"))
    assert payload["ok"] is True
    assert payload["output_path"] == str(output_path.resolve())
    assert saved["sheets"][0]["preview"][0][0] == "hello xlsx"


def test_agent_config_command_outputs_json(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)

    result = runner.invoke(app, ["agent-config"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "agent-config"
    assert payload["data"]["runtime"]["package"] == "docrt"
    assert "AGENTS.md" in payload["data"]["agents_md"]


def test_explain_task_command_outputs_json(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    task_path = tmp_path / "task.json"
    task_path.write_text(
        json.dumps({"task": "read-docx", "input": "sample.docx", "dry_run": True}),
        encoding="utf-8",
    )

    result = runner.invoke(app, ["explain-task", str(task_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "explain-task"
    assert payload["data"]["reads"] == ["sample.docx"]


def test_batch_fingerprint_command_outputs_json(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    xlsx_path = tmp_path / "sample.xlsx"
    document = Document()
    document.add_paragraph("hello docx")
    document.save(docx_path)
    workbook = openpyxl.Workbook()
    workbook.save(xlsx_path)

    result = runner.invoke(app, ["batch-fingerprint", str(docx_path), str(xlsx_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "batch-fingerprint"
    assert payload["data"]["count"] == 2
    assert payload["data"]["backend"] in {"python", "rust"}


def test_batch_read_isolates_per_file_errors(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    missing_path = tmp_path / "missing.pdf"
    document = Document()
    document.add_paragraph("hello docx")
    document.save(docx_path)

    result = runner.invoke(app, ["batch-read", str(docx_path), str(missing_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["data"]["count"] == 2
    assert payload["data"]["success_count"] == 1
    assert payload["data"]["failed_count"] == 1
    assert payload["data"]["results"][1]["ok"] is False
    assert payload["data"]["results"][1]["error"]["error_code"] == "FILE_NOT_FOUND"
    assert str(tmp_path) not in payload["data"]["results"][1]["error"]["error_message"]


def test_batch_inspect_uses_structural_inspect_results(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("hello docx")
    document.save(docx_path)

    result = runner.invoke(app, ["batch-inspect", str(docx_path)])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "batch-inspect"
    assert payload["backend"] == "inspect"
    assert payload["data"]["success_count"] == 1
    item = payload["data"]["results"][0]
    assert item["ok"] is True
    assert item["result"]["paragraph_count"] == 1
    assert "content_blocks" not in item["result"]


def test_clean_command_omits_file_list_by_default(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    log_path = tmp_path / "logs" / "old.jsonl"
    log_path.parent.mkdir()
    log_path.write_text("{}", encoding="utf-8")

    result = runner.invoke(app, ["clean", "--logs"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["data"]["files"] == []
    assert payload["data"]["files_omitted"] >= 1


def test_clean_retention_command_uses_default_targets(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    log_path = tmp_path / "logs" / "old.jsonl"
    log_path.parent.mkdir()
    log_path.write_text("{}", encoding="utf-8")

    result = runner.invoke(app, ["clean", "--retention", "--verbose"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["data"]["retention"] is True
    assert payload["data"]["selected_targets"] == ["logs", "diagnostics", "cache"]


def test_analyze_logs_command_outputs_recommendations(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    error_log = tmp_path / "logs" / "errors" / "2026-06-06.error.jsonl"
    error_log.parent.mkdir(parents=True)
    error_log.write_text(
        json.dumps(
            {
                "timestamp": "2026-06-06T00:00:00.000Z",
                "operation": "patch-docx",
                "module": "docrt.patch_ops",
                "error_code": "VALIDATION_FAILED",
                "exception_type": "ValidationError",
                "message": "expected_text mismatch",
            }
        )
        + "\n",
        encoding="utf-8",
    )

    result = runner.invoke(app, ["analyze-logs", "--days", "30"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "analyze-logs"
    assert payload["data"]["issue_count"] == 1
    assert payload["data"]["recommendations"]


def test_maintenance_command_writes_state(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)

    result = runner.invoke(app, ["maintenance"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "maintenance"
    assert Path(payload["data"]["state_paths"]["runtime_state"]).exists()


def test_repair_plan_command_writes_state(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)

    result = runner.invoke(app, ["repair-plan", "--days", "30"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["operation"] == "repair-plan"
    assert Path(payload["data"]["state_path"]).exists()


def test_job_start_rejects_unsupported_task(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)

    result = runner.invoke(app, ["job-start", "render-pdf"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["data"]["started"] is False
    assert "maintenance" in payload["data"]["supported_tasks"]
    assert "repair-plan" in payload["data"]["supported_tasks"]
    assert "clean-retention" in payload["data"]["supported_tasks"]


def test_job_start_clean_retention_command(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)

    result = runner.invoke(app, ["job-start", "clean-retention"])

    assert result.exit_code == 0
    payload = json.loads(result.stdout)
    assert payload["ok"] is True
    assert payload["data"]["started"] is True
    assert payload["data"]["task"] == "clean-retention"
    final = _wait_for_cli_job(Path(payload["data"]["status_path"]))
    assert final["status"] == "succeeded"


def _wait_for_cli_job(status_path: Path) -> dict[str, object]:
    deadline = time.monotonic() + 10
    while time.monotonic() < deadline:
        if status_path.exists():
            payload = json.loads(status_path.read_text(encoding="utf-8"))
            if payload.get("status") in {"succeeded", "failed"}:
                return payload
        time.sleep(0.1)
    return json.loads(status_path.read_text(encoding="utf-8"))
