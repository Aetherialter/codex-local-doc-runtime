import json
from pathlib import Path

import fitz
import openpyxl
import pytest
from docx import Document

from docrt.config import Config
from docrt.models import ErrorCode
from docrt.patch_ops import patch_docx, patch_xlsx
from docrt.paths import ValidationError
from docrt.read_ops import read_docx, read_pdf, read_xlsx
from docrt.task_ops import explain_task_manifest, run_task_manifest
from docrt.verify_ops import verify_docx, verify_xlsx


def _make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Original Heading", level=1)
    document.add_paragraph("hello docx")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Draft"
    document.save(path)


def _make_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet.append(["Field", "Value"])
    sheet.append(["Status", "Draft"])
    workbook.save(path)


def _make_pdf(path: Path) -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "hello pdf")
    document.save(path)
    document.close()


def test_read_docx_protocol(tmp_path: Path):
    path = tmp_path / "sample.docx"
    _make_docx(path)

    result = read_docx(path)

    assert result["document_type"] == "docx"
    assert result["metadata"]["paragraph_count"] == 2
    assert result["content_blocks"][0]["type"] == "heading"
    assert result["content_blocks"][0]["style"] == "Heading 1"
    assert result["content_blocks"][0]["is_heading"] is True
    assert result["tables"][0]["rows"][1][1]["text"] == "Draft"


def test_read_pdf_protocol(tmp_path: Path):
    path = tmp_path / "sample.pdf"
    _make_pdf(path)

    result = read_pdf(path)

    assert result["document_type"] == "pdf"
    assert result["metadata"]["page_count"] == 1
    assert result["metadata"]["has_text_layer"] is True
    assert result["content_blocks"][0]["type"] == "page_text"


def test_read_xlsx_protocol(tmp_path: Path):
    path = tmp_path / "sample.xlsx"
    _make_xlsx(path)

    result = read_xlsx(path)

    assert result["document_type"] == "xlsx"
    assert result["metadata"]["sheet_count"] == 1
    assert result["content_blocks"][0]["location"]["sheet"] == "Summary"
    assert result["metadata"]["sheets"][0]["preview_truncated"] is False


def test_read_xlsx_reports_preview_truncation(tmp_path: Path):
    path = tmp_path / "large.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Large"
    sheet.cell(25, 25).value = "outside preview"
    workbook.save(path)

    result = read_xlsx(path)

    sheet_info = result["metadata"]["sheets"][0]
    assert sheet_info["preview_range"] == "A1:T20"
    assert sheet_info["preview_truncated"] is True
    assert result["warnings"] == ["Sheet 'Large' preview is truncated to 20 rows x 20 columns."]


def test_patch_docx_and_verify(tmp_path: Path):
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    _make_docx(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "docx",
                "operations": [
                    {"type": "replace_text", "find": "hello", "replace": "patched"},
                    {
                        "type": "replace_table_cell",
                        "table_index": 0,
                        "row_index": 1,
                        "column_index": 1,
                        "text": "Ready",
                    },
                ],
            }
        ),
        encoding="utf-8",
    )

    patched = patch_docx(input_path, patch_path, output_path)
    verification = verify_docx(input_path, output_path)
    read_back = read_docx(output_path)

    assert patched["patch_summary"]["operation_count"] == 2
    assert patched["patch_summary"]["applied_count"] == 2
    assert verification["changed"] is True
    assert verification["changed_blocks"]
    assert "patched docx" in read_back["content_blocks"][1]["text"]
    assert read_back["tables"][0]["rows"][1][1]["text"] == "Ready"


def test_patch_docx_replace_heading_by_text_and_style(tmp_path: Path):
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    _make_docx(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "docx",
                "operations": [
                    {
                        "type": "replace_heading",
                        "heading_text": "Original",
                        "heading_style": "Heading 1",
                        "match": "contains",
                        "text": "Reviewed Heading",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    dry_run = patch_docx(input_path, patch_path, output_path, dry_run=True)
    patched = patch_docx(input_path, patch_path, output_path)
    read_back = read_docx(output_path)

    assert dry_run["patch_summary"]["planned_count"] == 1
    assert dry_run["patch_summary"]["applied_count"] == 0
    assert patched["patch_summary"]["applied_count"] == 1
    assert read_back["content_blocks"][0]["text"] == "Reviewed Heading"


def test_patch_docx_replace_heading_conflict_policy_replace(tmp_path: Path):
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    document = Document()
    document.add_heading("One", level=1)
    document.add_heading("Two", level=1)
    document.save(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "docx",
                "operations": [
                    {
                        "type": "replace_heading",
                        "heading_style": "Heading 1",
                        "text": "First Only",
                        "max_replacements": 1,
                        "conflict_policy": "replace",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    patched = patch_docx(input_path, patch_path, output_path)
    read_back = read_docx(output_path)

    assert patched["patch_summary"]["conflict_count"] == 1
    assert patched["patch_summary"]["skipped_count"] == 1
    assert read_back["content_blocks"][0]["text"] == "First Only"
    assert read_back["content_blocks"][1]["text"] == "Two"


def test_patch_xlsx_and_verify(tmp_path: Path):
    input_path = tmp_path / "sample.xlsx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.xlsx"
    _make_xlsx(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "xlsx",
                "operations": [
                    {"type": "set_cell", "sheet": "Summary", "cell": "B2", "value": "Ready"},
                    {"type": "add_sheet", "name": "Notes"},
                    {"type": "rename_sheet", "old_name": "Notes", "new_name": "Review"},
                ],
            }
        ),
        encoding="utf-8",
    )

    patched = patch_xlsx(input_path, patch_path, output_path)
    verification = verify_xlsx(input_path, output_path)
    read_back = read_xlsx(output_path)

    assert patched["patch_summary"]["operation_count"] == 3
    assert patched["patch_summary"]["applied_count"] == 3
    assert verification["changed"] is True
    assert verification["changed_cells"]
    assert any(block["text"] == "Ready" for block in read_back["content_blocks"])
    assert "Review" in [sheet["name"] for sheet in read_back["metadata"]["sheets"]]


def test_patch_docx_dry_run_does_not_write_output(tmp_path: Path):
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    _make_docx(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "docx",
                "operations": [
                    {
                        "type": "replace_paragraph",
                        "paragraph_index": 1,
                        "expected_text": "hello docx",
                        "text": "planned docx",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = patch_docx(input_path, patch_path, output_path, dry_run=True)

    assert result["dry_run"] is True
    assert result["verification"] is None
    assert result["patch_summary"]["planned_count"] == 1
    assert result["patch_summary"]["applied_count"] == 0
    assert result["patch_summary"]["planned_changes"]
    assert not output_path.exists()


def test_patch_xlsx_expected_value_mismatch_fails(tmp_path: Path):
    input_path = tmp_path / "sample.xlsx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.xlsx"
    _make_xlsx(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "xlsx",
                "operations": [
                    {
                        "type": "set_cell",
                        "sheet": "Summary",
                        "cell": "B2",
                        "expected_value": "Published",
                        "value": "Ready",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    with pytest.raises(ValidationError) as exc_info:
        patch_xlsx(input_path, patch_path, output_path)

    assert exc_info.value.error_code == ErrorCode.VALIDATION_FAILED


def test_patch_xlsx_preserves_existing_cell_style(tmp_path: Path):
    input_path = tmp_path / "sample.xlsx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet["B2"] = 12
    sheet["B2"].number_format = "0.00"
    sheet["B2"].font = openpyxl.styles.Font(bold=True)
    sheet["B2"].fill = openpyxl.styles.PatternFill("solid", fgColor="FFFF00")
    workbook.save(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "xlsx",
                "operations": [
                    {
                        "type": "set_cell",
                        "sheet": "Summary",
                        "cell": "B2",
                        "expected_value": 12,
                        "value": 13,
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = patch_xlsx(input_path, patch_path, output_path)
    patched = openpyxl.load_workbook(output_path)
    try:
        cell = patched["Summary"]["B2"]
        assert cell.value == 13
        assert cell.number_format == "0.00"
        assert cell.font.bold is True
        assert cell.fill.fgColor.rgb == "00FFFF00"
        assert (
            result["patch_summary"]["changes"][0]["format_preservation"]
            == "preserve_existing_cell_style"
        )
    finally:
        patched.close()


def test_patch_rejects_invalid_json(tmp_path: Path):
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    _make_docx(input_path)
    patch_path.write_text("{", encoding="utf-8")

    with pytest.raises(ValidationError) as exc_info:
        patch_docx(input_path, patch_path, output_path)

    assert exc_info.value.error_code == ErrorCode.VALIDATION_FAILED


def test_run_task_dry_run(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    manifest_path.write_text(
        json.dumps(
            {
                "task": "read-docx",
                "input": "examples/fixtures/sample.docx",
                "dry_run": True,
            }
        ),
        encoding="utf-8",
    )

    result = run_task_manifest(manifest_path, Config.load(project_root=tmp_path), "test-run")

    assert result["dry_run"] is True
    assert result["would_execute"]["task"] == "read-docx"


def test_run_task_search_pdf_writes_output(tmp_path: Path):
    input_path = tmp_path / "sample.pdf"
    output_path = tmp_path / "sample.pdf.search.json"
    manifest_path = tmp_path / "task.json"
    _make_pdf(input_path)
    manifest_path.write_text(
        json.dumps(
            {
                "task": "search-pdf",
                "input": str(input_path),
                "query": "hello",
                "output": str(output_path),
            }
        ),
        encoding="utf-8",
    )

    result = run_task_manifest(manifest_path, Config.load(project_root=tmp_path), "test-run")

    assert result["result"]["count"] == 1
    assert output_path.exists()


def test_run_task_search_pdf_pages_option(tmp_path: Path):
    pdf_path = tmp_path / "sample.pdf"
    output_path = tmp_path / "search.json"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "first target")
    page = document.new_page()
    page.insert_text((72, 72), "second target")
    document.save(pdf_path)
    document.close()
    task_path = tmp_path / "task.json"
    task_path.write_text(
        json.dumps(
            {
                "task": "search-pdf",
                "input": str(pdf_path),
                "query": "target",
                "pages": "2",
                "output": str(output_path),
            }
        ),
        encoding="utf-8",
    )

    result = run_task_manifest(task_path, Config.load(project_root=tmp_path), "test-run")

    assert result["result"]["count"] == 1
    assert result["result"]["page_selection"]["selected_pages"] == [2]


def test_run_task_multi_step_with_reference(tmp_path: Path):
    input_path = tmp_path / "sample.xlsx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.xlsx"
    manifest_path = tmp_path / "task.json"
    _make_xlsx(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "xlsx",
                "operations": [
                    {
                        "type": "set_cell",
                        "sheet": "Summary",
                        "cell": "B2",
                        "expected_value": "Draft",
                        "value": "Ready",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    manifest_path.write_text(
        json.dumps(
            {
                "stop_on_error": True,
                "tasks": [
                    {
                        "id": "patch",
                        "task": "patch-xlsx",
                        "input": str(input_path),
                        "patch": str(patch_path),
                        "output": str(output_path),
                    },
                    {
                        "id": "verify",
                        "task": "verify-xlsx",
                        "before": str(input_path),
                        "after": "${steps.patch.output_path}",
                    },
                ],
            }
        ),
        encoding="utf-8",
    )

    result = run_task_manifest(manifest_path, Config.load(project_root=tmp_path), "test-run")

    assert result["failed_count"] == 0
    assert result["success_count"] == 2
    assert result["steps"][1]["result"]["changed"] is True


def test_run_task_rejects_unknown_task(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    manifest_path.write_text(json.dumps({"task": "unknown", "input": "x"}), encoding="utf-8")

    with pytest.raises(ValidationError) as exc_info:
        run_task_manifest(manifest_path, Config.load(project_root=tmp_path), "test-run")

    assert exc_info.value.error_code == ErrorCode.VALIDATION_FAILED


def test_run_task_multi_step_error_contains_structured_error(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    manifest_path.write_text(
        json.dumps(
            {
                "stop_on_error": True,
                "tasks": [
                    {
                        "id": "read",
                        "task": "read-docx",
                        "input": str(tmp_path / "missing.docx"),
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = run_task_manifest(manifest_path, Config.load(project_root=tmp_path), "test-run")

    assert result["failed_count"] == 1
    assert result["steps"][0]["error"]["error_code"] == ErrorCode.FILE_NOT_FOUND.value
    assert result["steps"][0]["error"]["recovery_actions"]


def test_explain_task_manifest_reports_agent_effects(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    input_path = tmp_path / "sample.xlsx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.xlsx"
    manifest_path.write_text(
        json.dumps(
            {
                "stop_on_error": True,
                "tasks": [
                    {
                        "id": "patch",
                        "task": "patch-xlsx",
                        "input": str(input_path),
                        "patch": str(patch_path),
                        "output": str(output_path),
                        "dry_run": True,
                    },
                    {
                        "id": "verify",
                        "task": "verify-xlsx",
                        "before": str(input_path),
                        "after": "${steps.patch.output_path}",
                    },
                ],
            }
        ),
        encoding="utf-8",
    )

    result = explain_task_manifest(manifest_path)

    assert result["task_count"] == 2
    assert str(input_path) in result["reads"]
    assert "${steps.patch.output_path}" not in result["reads"]
    assert str(output_path) in result["writes"]
    assert str(output_path) in result["steps"][1]["reads"]
    assert str(patch_path) in result["patches"]
    assert any(
        item["input"] == str(input_path) and item["exists"] is False
        for item in result["path_resolution"]
    )
    assert result["requires_office_com"] is True
    assert result["supports_dry_run"] is True
    assert result["supports_native_dry_run"] is False
    assert result["dry_run_mode"] == "execute"
    assert result["steps"][0]["supports_native_dry_run"] is True
    assert result["steps"][0]["dry_run_mode"] == "native"
    assert result["steps"][1]["supports_native_dry_run"] is False
    assert result["steps"][1]["dry_run_mode"] == "plan_only"


def test_explain_task_manifest_reports_office_requirement(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    manifest_path.write_text(
        json.dumps({"task": "docx-to-pdf", "input": "sample.docx", "output": "sample.pdf"}),
        encoding="utf-8",
    )

    result = explain_task_manifest(manifest_path)

    assert result["requires_office_com"] is True
    assert result["generates"] == ["sample.pdf"]


def test_explain_task_manifest_reports_read_requires_office(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    manifest_path.write_text(
        json.dumps({"task": "read-pdf", "input": "sample.pdf"}),
        encoding="utf-8",
    )

    result = explain_task_manifest(manifest_path)

    assert result["requires_office_com"] is True
    assert result["steps"][0]["requires_office_com"] is True


def test_explain_task_manifest_reports_pdf_annotation_inputs(tmp_path: Path):
    manifest_path = tmp_path / "task.json"
    manifest_path.write_text(
        json.dumps(
            {
                "task": "annotate-pdf",
                "input": "sample.pdf",
                "annotations": "annotations.json",
                "output": "annotated.pdf",
            }
        ),
        encoding="utf-8",
    )

    result = explain_task_manifest(manifest_path)

    assert result["requires_office_com"] is True
    assert result["reads"] == ["sample.pdf", "annotations.json"]
    assert result["generates"] == ["annotated.pdf"]


def test_run_task_manifest_can_annotate_pdf(tmp_path: Path):
    input_path = tmp_path / "sample.pdf"
    annotations_path = tmp_path / "annotations.json"
    output_path = tmp_path / "annotated.pdf"
    manifest_path = tmp_path / "task.json"
    _make_pdf(input_path)
    annotations_path.write_text(
        json.dumps(
            {
                "annotations": [
                    {"type": "text_note", "page_number": 1, "point": [72, 72], "text": "note"}
                ]
            }
        ),
        encoding="utf-8",
    )
    manifest_path.write_text(
        json.dumps(
            {
                "task": "annotate-pdf",
                "input": str(input_path),
                "annotations": str(annotations_path),
                "output": str(output_path),
            }
        ),
        encoding="utf-8",
    )

    result = run_task_manifest(manifest_path, Config.load(project_root=tmp_path), "run")

    assert result["result"]["annotation_count"] == 1
    assert output_path.exists()
