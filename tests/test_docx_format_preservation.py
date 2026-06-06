from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from docrt.patch_ops import patch_docx


def test_replace_text_preserves_matching_run_format(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("hello")
    run.bold = True
    document.save(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "docx",
                "operations": [{"type": "replace_text", "find": "hello", "replace": "ready"}],
            }
        ),
        encoding="utf-8",
    )

    result = patch_docx(input_path, patch_path, output_path)
    patched = Document(str(output_path))

    assert result["patch_summary"]["changes"][0]["format_strategies"] == ["run_text"]
    assert patched.paragraphs[0].runs[0].text == "ready"
    assert patched.paragraphs[0].runs[0].bold is True


def test_replace_paragraph_preserves_style_and_first_run_format(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    patch_path = tmp_path / "patch.json"
    output_path = tmp_path / "patched.docx"
    document = Document()
    paragraph = document.add_paragraph(style="Heading 1")
    first = paragraph.add_run("Old")
    first.bold = True
    second = paragraph.add_run(" text")
    second.italic = True
    document.save(input_path)
    patch_path.write_text(
        json.dumps(
            {
                "document_type": "docx",
                "operations": [
                    {
                        "type": "replace_paragraph",
                        "paragraph_index": 0,
                        "expected_text": "Old text",
                        "text": "New heading",
                    }
                ],
            }
        ),
        encoding="utf-8",
    )

    result = patch_docx(input_path, patch_path, output_path)
    patched = Document(str(output_path))

    assert result["patch_summary"]["changes"][0]["format_preservation"] == "preserve_first_run"
    assert patched.paragraphs[0].style.name == "Heading 1"
    assert patched.paragraphs[0].runs[0].text == "New heading"
    assert patched.paragraphs[0].runs[0].bold is True
