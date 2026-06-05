from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

from docrt.core_bridge import validate_basic_json_object
from docrt.docx_styles import is_heading_style, paragraph_style_name
from docrt.models import ErrorCode
from docrt.paths import (
    ValidationError,
    ensure_output_not_locked,
    ensure_unlocked_for_read,
    validate_input_path,
    validate_output_path,
)
from docrt.read_ops import read_docx, read_xlsx


def patch_docx(
    input_path: str | Path,
    patch_path: str | Path,
    output_path: str | Path,
    *,
    dry_run: bool = False,
) -> dict[str, object]:
    source = validate_input_path(input_path, {".docx"})
    patch_file = validate_input_path(patch_path, {".json"})
    target = validate_output_path(output_path)
    _ensure_distinct_output(source, target)
    ensure_unlocked_for_read(source)
    ensure_unlocked_for_read(patch_file)
    if not dry_run:
        ensure_output_not_locked(target)
    patch = _load_patch(patch_file, expected_document_type="docx")

    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is unavailable") from exc

    document = Document(str(source))
    changes: list[dict[str, object]] = []
    skipped_count = 0
    conflicts: list[dict[str, object]] = []
    for operation in patch["operations"]:
        op_type = operation.get("type")
        if op_type == "replace_text":
            change = _docx_replace_text(document, operation, dry_run=dry_run)
        elif op_type == "replace_paragraph":
            change = _docx_replace_paragraph(document, operation, dry_run=dry_run)
        elif op_type == "replace_heading":
            change = _docx_replace_heading(document, operation, dry_run=dry_run)
        elif op_type == "replace_table_cell":
            change = _docx_replace_table_cell(document, operation, dry_run=dry_run)
        else:
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED, f"Unsupported DOCX patch op: {op_type}"
            )
        skipped_count += int(change.get("skipped_count", 0))
        if change.get("conflict"):
            conflicts.append(change)
        changes.append(change)

    verification: dict[str, object] | None = None
    if not dry_run:
        shutil.copyfile(source, target)
        document.save(str(target))
        verification = read_docx(target)

    return {
        "input_path": str(source),
        "patch_path": str(patch_file),
        "output_path": str(target),
        "dry_run": dry_run,
        "patch_summary": _patch_summary(changes, skipped_count, conflicts),
        "verification": _verification_summary(verification),
    }


def patch_xlsx(
    input_path: str | Path,
    patch_path: str | Path,
    output_path: str | Path,
    *,
    dry_run: bool = False,
) -> dict[str, object]:
    source = validate_input_path(input_path, {".xlsx"})
    patch_file = validate_input_path(patch_path, {".json"})
    target = validate_output_path(output_path)
    _ensure_distinct_output(source, target)
    ensure_unlocked_for_read(source)
    ensure_unlocked_for_read(patch_file)
    if not dry_run:
        ensure_output_not_locked(target)
    patch = _load_patch(patch_file, expected_document_type="xlsx")

    try:
        import openpyxl
    except Exception as exc:
        raise RuntimeError("openpyxl is unavailable") from exc

    workbook = openpyxl.load_workbook(str(source))
    changes: list[dict[str, object]] = []
    conflicts: list[dict[str, object]] = []
    try:
        for operation in patch["operations"]:
            op_type = operation.get("type")
            if op_type == "set_cell":
                operation_changes = [_xlsx_set_cell(workbook, operation, dry_run=dry_run)]
            elif op_type == "set_range_values":
                operation_changes = _xlsx_set_range_values(workbook, operation, dry_run=dry_run)
            elif op_type == "add_sheet":
                operation_changes = [_xlsx_add_sheet(workbook, operation, dry_run=dry_run)]
            elif op_type == "rename_sheet":
                operation_changes = [_xlsx_rename_sheet(workbook, operation, dry_run=dry_run)]
            else:
                raise ValidationError(
                    ErrorCode.VALIDATION_FAILED,
                    f"Unsupported XLSX patch op: {op_type}",
                )
            for change in operation_changes:
                if change.get("conflict"):
                    conflicts.append(change)
                changes.append(change)
        if not dry_run:
            shutil.copyfile(source, target)
            workbook.save(str(target))
    finally:
        workbook.close()

    verification = read_xlsx(target) if not dry_run else None
    return {
        "input_path": str(source),
        "patch_path": str(patch_file),
        "output_path": str(target),
        "dry_run": dry_run,
        "patch_summary": _patch_summary(changes, 0, conflicts),
        "verification": _verification_summary(verification),
    }


def _load_patch(path: Path, *, expected_document_type: str) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8")
    validate_basic_json_object(text)
    try:
        patch = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ValidationError(ErrorCode.VALIDATION_FAILED, f"Patch JSON is invalid: {exc}") from exc
    document_type = patch.get("document_type")
    if document_type != expected_document_type:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Patch document_type must be {expected_document_type!r}",
        )
    operations = patch.get("operations")
    if not isinstance(operations, list) or not operations:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED, "Patch operations must be a non-empty list"
        )
    for index, operation in enumerate(operations):
        if not isinstance(operation, dict) or not isinstance(operation.get("type"), str):
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED,
                f"Patch operation at index {index} must be an object with a string type",
            )
    return patch


def _ensure_distinct_output(source: Path, target: Path) -> None:
    if source == target:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "Patch commands require an output path different from the input path",
        )


def _patch_summary(
    changes: list[dict[str, object]],
    skipped_count: int,
    conflicts: list[dict[str, object]],
) -> dict[str, object]:
    return {
        "operation_count": len(changes),
        "planned_count": sum(int(change.get("planned_count", 1)) for change in changes),
        "applied_count": sum(int(change.get("applied_count", 1)) for change in changes),
        "skipped_count": skipped_count,
        "conflict_count": len(conflicts),
        "conflicts": conflicts,
        "planned_changes": changes,
        "changes": changes,
    }


def _verification_summary(verification: dict[str, object] | None) -> dict[str, object] | None:
    if verification is None:
        return None
    return {
        "document_type": verification["document_type"],
        "content_blocks": len(verification["content_blocks"]),
        "tables": len(verification["tables"]),
        "warnings": verification["warnings"],
    }


def _docx_replace_text(document, operation: dict[str, Any], *, dry_run: bool) -> dict[str, object]:
    find = _require_string(operation, "find")
    replace = _require_string(operation, "replace")
    scope = operation.get("scope", "all")
    if scope not in {"paragraphs", "tables", "all"}:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "scope must be one of: paragraphs, tables, all",
        )
    max_replacements = operation.get("max_replacements")
    if max_replacements is not None and (
        not isinstance(max_replacements, int) or max_replacements < 1
    ):
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "max_replacements must be a positive integer",
        )
    conflict_policy = operation.get("conflict_policy", "fail")
    if conflict_policy not in {"fail", "skip", "replace"}:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "conflict_policy must be one of: fail, skip, replace",
        )

    matches = _docx_text_matches(document, find, scope)
    total_matches = len(matches)
    skipped_count = 0
    conflict: str | None = None
    if max_replacements is not None and total_matches > max_replacements:
        conflict = "max_replacements_exceeded"
        if conflict_policy == "fail":
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED,
                f"replace_text matched {total_matches} locations, exceeding max_replacements "
                f"{max_replacements}",
            )
        if conflict_policy == "skip":
            skipped_count = total_matches
            matches = []
        else:
            skipped_count = total_matches - max_replacements
            matches = matches[:max_replacements]

    if not dry_run:
        _apply_docx_text_matches(document, matches, find, replace)
    result: dict[str, object] = {
        "type": "replace_text",
        "find": find,
        "replace": replace,
        "scope": scope,
        "planned_count": total_matches,
        "applied_count": 0 if dry_run else len(matches),
        "skipped_count": skipped_count,
        "locations": matches,
    }
    if conflict:
        result["conflict"] = conflict
    return result


def _docx_text_matches(document, find: str, scope: str) -> list[dict[str, object]]:
    matches: list[dict[str, object]] = []
    if scope in {"paragraphs", "all"}:
        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if find in paragraph.text:
                matches.append(
                    {
                        "type": "paragraph",
                        "paragraph_index": paragraph_index,
                        "old_text": paragraph.text,
                    }
                )
    if scope in {"tables", "all"}:
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    if find in cell.text:
                        matches.append(
                            {
                                "type": "table_cell",
                                "table_index": table_index,
                                "row_index": row_index,
                                "column_index": column_index,
                                "old_text": cell.text,
                            }
                        )
    return matches


def _apply_docx_text_matches(document, matches, find: str, replace: str) -> None:
    for match in matches:
        if match["type"] == "paragraph":
            paragraph = document.paragraphs[int(match["paragraph_index"])]
            paragraph.text = paragraph.text.replace(find, replace)
        else:
            cell = (
                document.tables[int(match["table_index"])]
                .rows[int(match["row_index"])]
                .cells[int(match["column_index"])]
            )
            cell.text = cell.text.replace(find, replace)


def _docx_replace_paragraph(
    document, operation: dict[str, Any], *, dry_run: bool
) -> dict[str, object]:
    index = _require_int(operation, "paragraph_index")
    text = _require_string(operation, "text")
    try:
        paragraph = document.paragraphs[index]
        old_text = paragraph.text
    except IndexError as exc:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED, f"Paragraph index out of range: {index}"
        ) from exc
    expected_text = operation.get("expected_text")
    if expected_text is not None and expected_text != old_text:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Paragraph {index} expected_text mismatch",
        )
    if not dry_run:
        paragraph.text = text
    return {
        "type": "replace_paragraph",
        "paragraph_index": index,
        "old_text": old_text,
        "text": text,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
    }


def _docx_replace_heading(
    document, operation: dict[str, Any], *, dry_run: bool
) -> dict[str, object]:
    replacement_text = _require_string(operation, "text")
    heading_text = operation.get("heading_text")
    heading_style = operation.get("heading_style")
    if heading_text is not None and not isinstance(heading_text, str):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "heading_text must be a string")
    if heading_style is not None and not isinstance(heading_style, str):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "heading_style must be a string")
    if heading_text is None and heading_style is None:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "replace_heading requires heading_text, heading_style, or both",
        )
    match_mode = operation.get("match", "exact")
    if match_mode not in {"exact", "contains"}:
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "match must be exact or contains")
    max_replacements = operation.get("max_replacements")
    if max_replacements is not None and (
        not isinstance(max_replacements, int) or max_replacements < 1
    ):
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "max_replacements must be a positive integer",
        )
    conflict_policy = operation.get("conflict_policy", "fail")
    if conflict_policy not in {"fail", "skip", "replace"}:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "conflict_policy must be one of: fail, skip, replace",
        )

    matches = _docx_heading_matches(
        document,
        heading_text=heading_text,
        heading_style=heading_style,
        match_mode=match_mode,
    )
    total_matches = len(matches)
    skipped_count = 0
    conflict: str | None = None
    if max_replacements is not None and total_matches > max_replacements:
        conflict = "max_replacements_exceeded"
        if conflict_policy == "fail":
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED,
                f"replace_heading matched {total_matches} locations, exceeding "
                f"max_replacements {max_replacements}",
            )
        if conflict_policy == "skip":
            skipped_count = total_matches
            matches = []
        else:
            skipped_count = total_matches - max_replacements
            matches = matches[:max_replacements]

    if not dry_run:
        for match in matches:
            document.paragraphs[int(match["paragraph_index"])].text = replacement_text

    result: dict[str, object] = {
        "type": "replace_heading",
        "heading_text": heading_text,
        "heading_style": heading_style,
        "match": match_mode,
        "text": replacement_text,
        "planned_count": total_matches,
        "applied_count": 0 if dry_run else len(matches),
        "skipped_count": skipped_count,
        "locations": matches,
    }
    if conflict:
        result["conflict"] = conflict
    return result


def _docx_heading_matches(
    document,
    *,
    heading_text: str | None,
    heading_style: str | None,
    match_mode: str,
) -> list[dict[str, object]]:
    matches: list[dict[str, object]] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        style_name = _paragraph_style_name(paragraph)
        if heading_style is not None:
            style_matches = style_name == heading_style
        else:
            style_matches = _is_heading_style(style_name)
        if not style_matches:
            continue
        if heading_text is not None and not _text_matches(paragraph.text, heading_text, match_mode):
            continue
        matches.append(
            {
                "type": "heading",
                "paragraph_index": paragraph_index,
                "old_text": paragraph.text,
                "style": style_name,
            }
        )
    return matches


def _paragraph_style_name(paragraph) -> str:
    return paragraph_style_name(paragraph)


def _is_heading_style(style_name: str) -> bool:
    return is_heading_style(style_name)


def _text_matches(value: str, expected: str, match_mode: str) -> bool:
    if match_mode == "contains":
        return expected in value
    return value == expected


def _docx_replace_table_cell(
    document, operation: dict[str, Any], *, dry_run: bool
) -> dict[str, object]:
    table_index = _require_int(operation, "table_index")
    row_index = _require_int(operation, "row_index")
    column_index = _require_int(operation, "column_index")
    text = _require_string(operation, "text")
    try:
        cell = document.tables[table_index].rows[row_index].cells[column_index]
    except IndexError as exc:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Table cell location out of range: {table_index}/{row_index}/{column_index}",
        ) from exc
    old_text = cell.text
    expected_text = operation.get("expected_text")
    if expected_text is not None and expected_text != old_text:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Table cell {table_index}/{row_index}/{column_index} expected_text mismatch",
        )
    if not dry_run:
        cell.text = text
    return {
        "type": "replace_table_cell",
        "table_index": table_index,
        "row_index": row_index,
        "column_index": column_index,
        "old_text": old_text,
        "text": text,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
    }


def _xlsx_set_cell(workbook, operation: dict[str, Any], *, dry_run: bool) -> dict[str, object]:
    sheet = _worksheet(workbook, _require_string(operation, "sheet"))
    cell = _require_string(operation, "cell")
    value = operation.get("value")
    old_value = sheet[cell].value
    expected_value = operation.get("expected_value")
    if "expected_value" in operation and expected_value != old_value:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Cell {sheet.title}!{cell} expected_value mismatch",
        )
    if not dry_run:
        sheet[cell] = value
    return {
        "type": "set_cell",
        "sheet": sheet.title,
        "cell": cell,
        "old_value": old_value,
        "value": value,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
    }


def _xlsx_set_range_values(
    workbook, operation: dict[str, Any], *, dry_run: bool
) -> list[dict[str, object]]:
    sheet = _worksheet(workbook, _require_string(operation, "sheet"))
    start_cell = _require_string(operation, "start_cell")
    values = operation.get("values")
    if (
        not isinstance(values, list)
        or not values
        or not all(isinstance(row, list) for row in values)
    ):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "values must be a non-empty 2D list")
    expected_values = operation.get("expected_values")
    if expected_values is not None and not (
        isinstance(expected_values, list) and all(isinstance(row, list) for row in expected_values)
    ):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "expected_values must be a 2D list")
    start = sheet[start_cell]
    changes = []
    for row_offset, row_values in enumerate(values):
        for column_offset, value in enumerate(row_values):
            cell = sheet.cell(start.row + row_offset, start.column + column_offset)
            old_value = cell.value
            if expected_values is not None:
                try:
                    expected_value = expected_values[row_offset][column_offset]
                except IndexError as exc:
                    raise ValidationError(
                        ErrorCode.VALIDATION_FAILED,
                        "expected_values shape must match values shape",
                    ) from exc
                if expected_value != old_value:
                    raise ValidationError(
                        ErrorCode.VALIDATION_FAILED,
                        f"Cell {sheet.title}!{cell.coordinate} expected_values mismatch",
                    )
            if not dry_run:
                cell.value = value
            changes.append(
                {
                    "type": "set_cell",
                    "sheet": sheet.title,
                    "cell": cell.coordinate,
                    "old_value": old_value,
                    "value": value,
                    "planned_count": 1,
                    "applied_count": 0 if dry_run else 1,
                }
            )
    return changes


def _xlsx_add_sheet(workbook, operation: dict[str, Any], *, dry_run: bool) -> dict[str, object]:
    name = _require_string(operation, "name")
    if name in workbook.sheetnames:
        raise ValidationError(ErrorCode.VALIDATION_FAILED, f"Sheet already exists: {name}")
    if not dry_run:
        workbook.create_sheet(name)
    return {
        "type": "add_sheet",
        "name": name,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
    }


def _xlsx_rename_sheet(workbook, operation: dict[str, Any], *, dry_run: bool) -> dict[str, object]:
    old_name = _require_string(operation, "old_name")
    new_name = _require_string(operation, "new_name")
    if new_name in workbook.sheetnames:
        raise ValidationError(ErrorCode.VALIDATION_FAILED, f"Sheet already exists: {new_name}")
    sheet = _worksheet(workbook, old_name)
    if not dry_run:
        sheet.title = new_name
    return {
        "type": "rename_sheet",
        "old_name": old_name,
        "new_name": new_name,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
    }


def _worksheet(workbook, name: str):
    if name not in workbook.sheetnames:
        raise ValidationError(ErrorCode.VALIDATION_FAILED, f"Sheet not found: {name}")
    return workbook[name]


def _require_string(operation: dict[str, Any], key: str) -> str:
    value = operation.get(key)
    if not isinstance(value, str):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, f"{key} must be a string")
    return value


def _require_int(operation: dict[str, Any], key: str) -> int:
    value = operation.get(key)
    if not isinstance(value, int):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, f"{key} must be an integer")
    return value
