from __future__ import annotations

import shutil
from pathlib import Path
from typing import Any

from docrt.docx_styles import is_heading_style, paragraph_style_name
from docrt.models import ErrorCode
from docrt.patch_common import (
    ensure_distinct_output,
    load_patch,
    patch_summary,
    require_int,
    require_string,
    verification_summary,
)
from docrt.paths import (
    ValidationError,
    ensure_output_not_locked,
    ensure_unlocked_for_read,
    validate_input_path,
    validate_output_path,
)
from docrt.read_ops import read_docx


def patch_docx(
    input_path: str | Path,
    patch_path: str | Path,
    output_path: str | Path,
    *,
    dry_run: bool = False,
) -> dict[str, object]:
    source = validate_input_path(input_path, {".docx"})
    patch_file = validate_input_path(patch_path, {".json"})
    target = validate_output_path(output_path)
    ensure_distinct_output(source, target)
    ensure_unlocked_for_read(source)
    ensure_unlocked_for_read(patch_file)
    if not dry_run:
        ensure_output_not_locked(target)
    patch = load_patch(patch_file, expected_document_type="docx")

    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is unavailable") from exc

    document = Document(str(source))
    changes: list[dict[str, object]] = []
    skipped_count = 0
    conflicts: list[dict[str, object]] = []
    for operation in patch["operations"]:
        op_type = operation.get("type")
        if op_type == "replace_text":
            change = _docx_replace_text(document, operation, dry_run=dry_run)
        elif op_type == "replace_paragraph":
            change = _docx_replace_paragraph(document, operation, dry_run=dry_run)
        elif op_type == "replace_heading":
            change = _docx_replace_heading(document, operation, dry_run=dry_run)
        elif op_type == "replace_table_cell":
            change = _docx_replace_table_cell(document, operation, dry_run=dry_run)
        else:
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED, f"Unsupported DOCX patch op: {op_type}"
            )
        skipped_count += int(change.get("skipped_count", 0))
        if change.get("conflict"):
            conflicts.append(change)
        changes.append(change)

    verification: dict[str, object] | None = None
    if not dry_run:
        shutil.copyfile(source, target)
        document.save(str(target))
        verification = read_docx(target)

    return {
        "input_path": str(source),
        "patch_path": str(patch_file),
        "output_path": str(target),
        "dry_run": dry_run,
        "patch_summary": patch_summary(changes, skipped_count, conflicts),
        "verification": verification_summary(verification),
    }


def _docx_replace_text(document, operation: dict[str, Any], *, dry_run: bool) -> dict[str, object]:
    find = require_string(operation, "find")
    replace = require_string(operation, "replace")
    scope = operation.get("scope", "all")
    if scope not in {"paragraphs", "tables", "all"}:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "scope must be one of: paragraphs, tables, all",
        )
    max_replacements = operation.get("max_replacements")
    if max_replacements is not None and (
        not isinstance(max_replacements, int) or max_replacements < 1
    ):
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "max_replacements must be a positive integer",
        )
    conflict_policy = operation.get("conflict_policy", "fail")
    if conflict_policy not in {"fail", "skip", "replace"}:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "conflict_policy must be one of: fail, skip, replace",
        )

    matches = _docx_text_matches(document, find, scope)
    total_matches = len(matches)
    skipped_count = 0
    conflict: str | None = None
    if max_replacements is not None and total_matches > max_replacements:
        conflict = "max_replacements_exceeded"
        if conflict_policy == "fail":
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED,
                f"replace_text matched {total_matches} locations, exceeding max_replacements "
                f"{max_replacements}",
            )
        if conflict_policy == "skip":
            skipped_count = total_matches
            matches = []
        else:
            skipped_count = total_matches - max_replacements
            matches = matches[:max_replacements]

    format_strategies: list[str] = []
    if not dry_run:
        format_strategies = _apply_docx_text_matches(document, matches, find, replace)
    result: dict[str, object] = {
        "type": "replace_text",
        "find": find,
        "replace": replace,
        "scope": scope,
        "planned_count": total_matches,
        "applied_count": 0 if dry_run else len(matches),
        "skipped_count": skipped_count,
        "locations": matches,
        "format_preservation": "run_aware",
    }
    if format_strategies:
        result["format_strategies"] = format_strategies
    if conflict:
        result["conflict"] = conflict
    return result


def _docx_text_matches(document, find: str, scope: str) -> list[dict[str, object]]:
    matches: list[dict[str, object]] = []
    if scope in {"paragraphs", "all"}:
        for paragraph_index, paragraph in enumerate(document.paragraphs):
            if find in paragraph.text:
                matches.append(
                    {
                        "type": "paragraph",
                        "paragraph_index": paragraph_index,
                        "old_text": paragraph.text,
                    }
                )
    if scope in {"tables", "all"}:
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    if find in cell.text:
                        matches.append(
                            {
                                "type": "table_cell",
                                "table_index": table_index,
                                "row_index": row_index,
                                "column_index": column_index,
                                "old_text": cell.text,
                            }
                        )
    return matches


def _apply_docx_text_matches(document, matches, find: str, replace: str) -> list[str]:
    strategies: list[str] = []
    for match in matches:
        if match["type"] == "paragraph":
            paragraph = document.paragraphs[int(match["paragraph_index"])]
            strategies.append(_replace_text_in_paragraph(paragraph, find, replace))
        else:
            cell = (
                document.tables[int(match["table_index"])]
                .rows[int(match["row_index"])]
                .cells[int(match["column_index"])]
            )
            strategies.append(_replace_text_in_cell(cell, find, replace))
    return strategies


def _docx_replace_paragraph(
    document, operation: dict[str, Any], *, dry_run: bool
) -> dict[str, object]:
    index = require_int(operation, "paragraph_index")
    text = require_string(operation, "text")
    try:
        paragraph = document.paragraphs[index]
        old_text = paragraph.text
    except IndexError as exc:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED, f"Paragraph index out of range: {index}"
        ) from exc
    expected_text = operation.get("expected_text")
    if expected_text is not None and expected_text != old_text:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Paragraph {index} expected_text mismatch",
        )
    format_strategy = _planned_set_text_strategy(paragraph)
    if not dry_run:
        format_strategy = _set_paragraph_text_preserve_first_run(paragraph, text)
    return {
        "type": "replace_paragraph",
        "paragraph_index": index,
        "old_text": old_text,
        "text": text,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
        "format_preservation": format_strategy,
    }


def _docx_replace_heading(
    document, operation: dict[str, Any], *, dry_run: bool
) -> dict[str, object]:
    replacement_text = require_string(operation, "text")
    heading_text = operation.get("heading_text")
    heading_style = operation.get("heading_style")
    if heading_text is not None and not isinstance(heading_text, str):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "heading_text must be a string")
    if heading_style is not None and not isinstance(heading_style, str):
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "heading_style must be a string")
    if heading_text is None and heading_style is None:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "replace_heading requires heading_text, heading_style, or both",
        )
    match_mode = operation.get("match", "exact")
    if match_mode not in {"exact", "contains"}:
        raise ValidationError(ErrorCode.VALIDATION_FAILED, "match must be exact or contains")
    max_replacements = operation.get("max_replacements")
    if max_replacements is not None and (
        not isinstance(max_replacements, int) or max_replacements < 1
    ):
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "max_replacements must be a positive integer",
        )
    conflict_policy = operation.get("conflict_policy", "fail")
    if conflict_policy not in {"fail", "skip", "replace"}:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            "conflict_policy must be one of: fail, skip, replace",
        )

    matches = _docx_heading_matches(
        document,
        heading_text=heading_text,
        heading_style=heading_style,
        match_mode=match_mode,
    )
    total_matches = len(matches)
    skipped_count = 0
    conflict: str | None = None
    if max_replacements is not None and total_matches > max_replacements:
        conflict = "max_replacements_exceeded"
        if conflict_policy == "fail":
            raise ValidationError(
                ErrorCode.VALIDATION_FAILED,
                f"replace_heading matched {total_matches} locations, exceeding "
                f"max_replacements {max_replacements}",
            )
        if conflict_policy == "skip":
            skipped_count = total_matches
            matches = []
        else:
            skipped_count = total_matches - max_replacements
            matches = matches[:max_replacements]

    format_strategies: list[str] = []
    if not dry_run:
        for match in matches:
            paragraph = document.paragraphs[int(match["paragraph_index"])]
            format_strategies.append(
                _set_paragraph_text_preserve_first_run(paragraph, replacement_text)
            )

    result: dict[str, object] = {
        "type": "replace_heading",
        "heading_text": heading_text,
        "heading_style": heading_style,
        "match": match_mode,
        "text": replacement_text,
        "planned_count": total_matches,
        "applied_count": 0 if dry_run else len(matches),
        "skipped_count": skipped_count,
        "locations": matches,
        "format_preservation": "preserve_paragraph_style_and_first_run",
    }
    if format_strategies:
        result["format_strategies"] = format_strategies
    if conflict:
        result["conflict"] = conflict
    return result


def _docx_heading_matches(
    document,
    *,
    heading_text: str | None,
    heading_style: str | None,
    match_mode: str,
) -> list[dict[str, object]]:
    matches: list[dict[str, object]] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph_style_name(paragraph)
        if heading_style is not None:
            style_matches = style_name == heading_style
        else:
            style_matches = is_heading_style(style_name)
        if not style_matches:
            continue
        if heading_text is not None and not _text_matches(paragraph.text, heading_text, match_mode):
            continue
        matches.append(
            {
                "type": "heading",
                "paragraph_index": paragraph_index,
                "old_text": paragraph.text,
                "style": style_name,
            }
        )
    return matches


def _text_matches(value: str, expected: str, match_mode: str) -> bool:
    if match_mode == "contains":
        return expected in value
    return value == expected


def _docx_replace_table_cell(
    document, operation: dict[str, Any], *, dry_run: bool
) -> dict[str, object]:
    table_index = require_int(operation, "table_index")
    row_index = require_int(operation, "row_index")
    column_index = require_int(operation, "column_index")
    text = require_string(operation, "text")
    try:
        cell = document.tables[table_index].rows[row_index].cells[column_index]
    except IndexError as exc:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Table cell location out of range: {table_index}/{row_index}/{column_index}",
        ) from exc
    old_text = cell.text
    expected_text = operation.get("expected_text")
    if expected_text is not None and expected_text != old_text:
        raise ValidationError(
            ErrorCode.VALIDATION_FAILED,
            f"Table cell {table_index}/{row_index}/{column_index} expected_text mismatch",
        )
    format_strategy = _planned_cell_strategy(cell)
    if not dry_run:
        format_strategy = _set_cell_text_preserve_first_run(cell, text)
    return {
        "type": "replace_table_cell",
        "table_index": table_index,
        "row_index": row_index,
        "column_index": column_index,
        "old_text": old_text,
        "text": text,
        "planned_count": 1,
        "applied_count": 0 if dry_run else 1,
        "format_preservation": format_strategy,
    }


def _replace_text_in_cell(cell, find: str, replace: str) -> str:
    strategies = [
        _replace_text_in_paragraph(paragraph, find, replace)
        for paragraph in cell.paragraphs
        if find in paragraph.text
    ]
    if strategies:
        return ",".join(sorted(set(strategies)))
    return _set_cell_text_preserve_first_run(cell, cell.text.replace(find, replace))


def _replace_text_in_paragraph(paragraph, find: str, replace: str) -> str:
    replaced_in_run = False
    for run in paragraph.runs:
        if find in run.text:
            run.text = run.text.replace(find, replace)
            replaced_in_run = True
    if replaced_in_run:
        return "run_text"
    new_text = paragraph.text.replace(find, replace)
    return _set_paragraph_text_preserve_first_run(paragraph, new_text)


def _set_cell_text_preserve_first_run(cell, text: str) -> str:
    if not cell.paragraphs:
        cell.text = text
        return "cell_text"
    strategy = _set_paragraph_text_preserve_first_run(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        _set_paragraph_text_preserve_first_run(paragraph, "")
    return f"cell_{strategy}"


def _set_paragraph_text_preserve_first_run(paragraph, text: str) -> str:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
        return "preserve_first_run"
    paragraph.text = text
    return "paragraph_text"


def _planned_set_text_strategy(paragraph) -> str:
    return "preserve_first_run" if paragraph.runs else "paragraph_text"


def _planned_cell_strategy(cell) -> str:
    if cell.paragraphs and cell.paragraphs[0].runs:
        return "cell_preserve_first_run"
    return "cell_text"
