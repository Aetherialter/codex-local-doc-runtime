from __future__ import annotations

from pathlib import Path

from docrt.paths import ensure_unlocked_for_read, validate_input_path


def inspect_docx(path: str | Path) -> dict[str, object]:
    input_path = validate_input_path(path, {".docx"})
    ensure_unlocked_for_read(input_path)
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is unavailable") from exc

    document = Document(str(input_path))
    paragraphs = [p.text for p in document.paragraphs]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    return {
        "path": str(input_path),
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "paragraphs": paragraphs,
        "tables": tables,
    }
