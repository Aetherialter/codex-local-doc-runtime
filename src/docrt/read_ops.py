from __future__ import annotations

from pathlib import Path
from typing import Any

from docrt.docx_styles import is_heading_style, paragraph_style_name
from docrt.paths import ensure_unlocked_for_read, validate_input_path
from docrt.pdf_pages import page_selection_metadata, selected_page_indexes
from docrt.pdf_safety import ensure_pdf_not_encrypted, pdf_text_layer_warnings


def read_docx(path: str | Path) -> dict[str, object]:
    input_path = validate_input_path(path, {".docx"})
    ensure_unlocked_for_read(input_path)
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is unavailable") from exc

    document = Document(str(input_path))
    content_blocks: list[dict[str, object]] = []
    for index, paragraph in enumerate(document.paragraphs):
        style_name = paragraph_style_name(paragraph)
        content_blocks.append(
            {
                "type": "heading" if is_heading_style(style_name) else "paragraph",
                "text": paragraph.text,
                "location": {"paragraph_index": index},
                "style": style_name,
                "is_heading": is_heading_style(style_name),
            }
        )

    tables: list[dict[str, object]] = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row_index, row in enumerate(table.rows):
            cells = []
            for column_index, cell in enumerate(row.cells):
                text = cell.text
                cells.append(
                    {
                        "text": text,
                        "location": {
                            "table_index": table_index,
                            "row_index": row_index,
                            "column_index": column_index,
                        },
                    }
                )
                content_blocks.append(
                    {
                        "type": "table_cell",
                        "text": text,
                        "location": {
                            "table_index": table_index,
                            "row_index": row_index,
                            "column_index": column_index,
                        },
                    }
                )
            rows.append(cells)
        tables.append({"table_index": table_index, "rows": rows})

    return {
        "document_type": "docx",
        "path": str(input_path),
        "metadata": {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
        "content_blocks": content_blocks,
        "tables": tables,
        "warnings": [],
    }


def read_pdf(path: str | Path, *, pages: str | None = None) -> dict[str, object]:
    input_path = validate_input_path(path, {".pdf"})
    ensure_unlocked_for_read(input_path)
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PyMuPDF is unavailable") from exc

    document = fitz.open(str(input_path))
    try:
        ensure_pdf_not_encrypted(document)
        content_blocks: list[dict[str, object]] = []
        page_summaries: list[dict[str, Any]] = []
        total_text_chars = 0
        selected_indexes = selected_page_indexes(document.page_count, pages)
        for index in selected_indexes:
            page = document[index]
            text = page.get_text("text")
            total_text_chars += len(text)
            rect = page.rect
            page_number = index + 1
            page_summaries.append(
                {
                    "page_number": page_number,
                    "width": rect.width,
                    "height": rect.height,
                    "text_chars": len(text),
                }
            )
            content_blocks.append(
                {
                    "type": "page_text",
                    "text": text,
                    "location": {
                        "page_number": page_number,
                        "width": rect.width,
                        "height": rect.height,
                    },
                }
            )
        warnings = pdf_text_layer_warnings(total_text_chars)
        return {
            "document_type": "pdf",
            "path": str(input_path),
            "metadata": {
                "page_count": document.page_count,
                "page_selection": page_selection_metadata(
                    document.page_count, selected_indexes, pages
                ),
                "pdf_metadata": dict(document.metadata or {}),
                "has_text_layer": total_text_chars > 0,
                "needs_ocr": total_text_chars == 0,
                "ocr_supported": False,
                "encryption": {
                    "is_encrypted": bool(getattr(document, "is_encrypted", False)),
                    "needs_pass": bool(getattr(document, "needs_pass", False)),
                },
                "pages": page_summaries,
            },
            "content_blocks": content_blocks,
            "tables": [],
            "warnings": warnings,
        }
    finally:
        document.close()


def read_xlsx(path: str | Path) -> dict[str, object]:
    input_path = validate_input_path(path, {".xlsx"})
    ensure_unlocked_for_read(input_path)
    try:
        import openpyxl
    except Exception as exc:
        raise RuntimeError("openpyxl is unavailable") from exc

    workbook = openpyxl.load_workbook(str(input_path), data_only=False, read_only=False)
    try:
        content_blocks: list[dict[str, object]] = []
        tables: list[dict[str, object]] = []
        sheets: list[dict[str, object]] = []
        warnings: list[str] = []
        for sheet in workbook.worksheets:
            rows = []
            actual_max_row = sheet.max_row or 0
            actual_max_col = sheet.max_column or 0
            max_row = min(actual_max_row, 20)
            max_col = min(actual_max_col, 20)
            truncated = actual_max_row > max_row or actual_max_col > max_col
            if truncated:
                warnings.append(
                    f"Sheet {sheet.title!r} preview is truncated to 20 rows x 20 columns."
                )
            for row in sheet.iter_rows(
                min_row=1,
                max_row=max_row,
                min_col=1,
                max_col=max_col,
            ):
                row_values = []
                for cell in row:
                    value = cell.value
                    row_values.append(value)
                    if value is not None:
                        content_blocks.append(
                            {
                                "type": "cell",
                                "text": str(value),
                                "location": {
                                    "sheet": sheet.title,
                                    "cell": cell.coordinate,
                                    "row": cell.row,
                                    "column": cell.column,
                                },
                            }
                        )
                rows.append(row_values)
            end_cell = sheet.cell(max_row, max_col).coordinate if max_row and max_col else "A1"
            tables.append({"sheet": sheet.title, "range": f"A1:{end_cell}", "rows": rows})
            sheets.append(
                {
                    "name": sheet.title,
                    "max_row": sheet.max_row,
                    "max_column": sheet.max_column,
                    "preview_range": f"A1:{end_cell}",
                    "preview_truncated": truncated,
                    "merged_cells": [str(rng) for rng in sheet.merged_cells.ranges],
                }
            )

        return {
            "document_type": "xlsx",
            "path": str(input_path),
            "metadata": {
                "sheet_count": len(workbook.worksheets),
                "sheets": sheets,
            },
            "content_blocks": content_blocks,
            "tables": tables,
            "warnings": warnings,
        }
    finally:
        workbook.close()
